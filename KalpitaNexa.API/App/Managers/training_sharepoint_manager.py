"""
Manager for Training module SharePoint operations via Microsoft Graph API.
Handles all file reads and writes — NO business logic here.

Place at: KalpitaNexa/KalpitaNexa.API/App/Managers/training_sharepoint_manager.py

Real SharePoint structure:
  Library : Documents
  Root    : HR/KT-Merin-27 june 2025/2025/1. Training/AI Training - Internal - Dev and DB/
  Per day : <Day Folder>/
                <Day X - transcript - Name>.docx   <- auto-detected
                <Day X - Training Video - Name>.mp4 <- auto-detected (Phase 2)
            <Day Folder>/output/
                summary.docx
                qa.docx
                processing_complete.json
"""

import logging
import json
import io
from datetime import datetime
from typing import Optional

import requests
from msal import ConfidentialClientApplication
from docx import Document

from .. import config

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

SHAREPOINT_HOST = "kalpitatechnologies0.sharepoint.com"
SHAREPOINT_SITE = "sites/KalpitaMasterFolder"

# Full folder path inside the Documents library leading to the Day folders
TRAINING_ROOT = (
    "HR/KT-Merin-27 june 2025/2025/1. Training"
    "/AI Training - Internal - Dev and DB"
)

GRAPH_BASE  = "https://graph.microsoft.com/v1.0"
GRAPH_SCOPE = ["https://graph.microsoft.com/.default"]


class TrainingSharePointManager:
    """
    All SharePoint I/O for the Training module.
    Uses Service Principal (client credentials) auth via MSAL.

    Key design: file names inside each Day folder are NOT hardcoded.
    The manager lists folder contents and detects:
      - Transcript : any .docx whose name contains 'transcript' (case-insensitive)
      - Video      : any .mp4 file (Phase 2)
    """

    def __init__(self):
        self._token: Optional[str]    = None #These are **cached values** — once we resolve them, we don't look them up again
        self._site_id: Optional[str]  = None
        self._drive_id: Optional[str] = None

    # ------------------------------------------------------------------
    # Auth
    # ------------------------------------------------------------------

    def _get_token(self) -> str:#Check if we already have a token → return it immediately
        """Acquire (or reuse) a Graph API access token via Service Principal."""
        if self._token:
            return self._token

        app = ConfidentialClientApplication(#Returns a token valid for ~1 hour
            client_id=config.SP_TRAINING_CLIENT_ID,
            client_credential=config.SP_TRAINING_CLIENT_SECRET,
            authority=(
                f"https://login.microsoftonline.com/{config.SP_TRAINING_TENANT_ID}"
            ),
        )
        result = app.acquire_token_for_client(scopes=GRAPH_SCOPE)

        if "access_token" not in result:
            raise RuntimeError(
                f"Failed to acquire Graph token: "
                f"{result.get('error_description', result)}"
            )

        self._token = result["access_token"]
        logger.info("Graph API token acquired for TrainingSharePointManager.")
        return self._token

    def _auth_headers(self) -> dict:
        return {"Authorization": f"Bearer {self._get_token()}"}

    # ------------------------------------------------------------------
    # Site / Drive resolution (cached after first call)
    # ------------------------------------------------------------------

    def _get_site_id(self) -> str:
        if self._site_id:
            return self._site_id

        url  = f"{GRAPH_BASE}/sites/{SHAREPOINT_HOST}:/{SHAREPOINT_SITE}"#Graph API needs the **site ID** (a GUID) to access SharePoint
        resp = requests.get(url, headers=self._auth_headers())
        resp.raise_for_status()#throws exception if HTTP status >= 400
        self._site_id = resp.json()["id"]
        logger.debug(f"Resolved site_id: {self._site_id}")
        return self._site_id

    def _get_drive_id(self) -> str:
        if self._drive_id:
            return self._drive_id

        site_id = self._get_site_id()
        url     = f"{GRAPH_BASE}/sites/{site_id}/drives"
        resp    = requests.get(url, headers=self._auth_headers())
        resp.raise_for_status()

        drives    = resp.json().get("value", [])
        doc_drive = next(
            (d for d in drives if d.get("name") == "Documents"),
            drives[0],
        )
        self._drive_id = doc_drive["id"]
        logger.debug(f"Resolved drive_id: {self._drive_id}")
        return self._drive_id

    # ------------------------------------------------------------------
    # Path helpers
    # ------------------------------------------------------------------

    def _item_url(self, relative_path: str) -> str:
        """Return the Graph URL for a drive item by its library-relative path."""
        drive_id = self._get_drive_id()
        return f"{GRAPH_BASE}/drives/{drive_id}/root:/{relative_path}"

    def _day_folder_path(self, day_folder_name: str) -> str:
        return f"{TRAINING_ROOT}/{day_folder_name}"

    def _output_folder_path(self, day_folder_name: str) -> str:
        return f"{self._day_folder_path(day_folder_name)}/output"

    # ------------------------------------------------------------------
    # Dynamic file detection
    # ------------------------------------------------------------------

    def _list_folder_files(self, folder_path: str) -> list:
        """
        Return a list of file-item dicts directly inside folder_path.
        Returns [] if folder does not exist.
        """
        drive_id = self._get_drive_id()
        url      = f"{GRAPH_BASE}/drives/{drive_id}/root:/{folder_path}:/children"
        resp     = requests.get(url, headers=self._auth_headers())

        if resp.status_code == 404:
            return []
        resp.raise_for_status()

        # Only files (items that have a "file" key), not subfolders
        return [i for i in resp.json().get("value", []) if "file" in i]

    def _find_transcript_name(self, day_folder_name: str) -> Optional[str]:
        """
        Scan the Day folder and return the filename of the transcript docx.
        Matches any .docx containing 'transcript' in its name (case-insensitive).
        """
        items = self._list_folder_files(self._day_folder_path(day_folder_name))
        for item in items:
            name = item["name"]
            if name.lower().endswith(".docx") and "transcript" in name.lower():
                logger.debug(f"Detected transcript: '{name}'")
                return name
        logger.warning(f"No transcript .docx found in: '{day_folder_name}'")
        return None

    def find_video_name(self, day_folder_name: str) -> Optional[str]:
        """
        Scan the Day folder and return the filename of the .mp4 video.
        Used by Phase 2.
        """
        items = self._list_folder_files(self._day_folder_path(day_folder_name))
        for item in items:
            name = item["name"]
            if name.lower().endswith(".mp4"):
                logger.debug(f"Detected video: '{name}'")
                return name
        logger.warning(f"No .mp4 found in: '{day_folder_name}'")
        return None

    # ------------------------------------------------------------------
    # Public: file existence checks
    # ------------------------------------------------------------------

    def file_exists(self, library_relative_path: str) -> bool:
        """Return True if a file exists at the given library-relative path."""
        try:
            resp = requests.get(
                self._item_url(library_relative_path),
                headers=self._auth_headers(),
            )
            return resp.status_code == 200
        except Exception as e:
            logger.warning(f"file_exists check failed for '{library_relative_path}': {e}")
            return False

    def output_file_exists(self, day_folder_name: str, file_name: str) -> bool:
        """Check if a file exists in the Day folder's output/ subfolder."""
        path = f"{self._output_folder_path(day_folder_name)}/{file_name}"
        return self.file_exists(path)

    # ------------------------------------------------------------------
    # Public: Read
    # ------------------------------------------------------------------

    def read_transcript(self, day_folder_name: str) -> str:
        """
        Auto-detect and download the transcript .docx from the Day folder.
        Returns full plain text with paragraphs joined by newlines.
        Raises FileNotFoundError if no transcript docx is found.
        """
        #We keep it in memory (BytesIO) for efficiency — no disk I/O overhead. The file is <1MB so memory isn't a concern."
        transcript_name = self._find_transcript_name(day_folder_name)
        if not transcript_name:
            raise FileNotFoundError(
                f"No transcript .docx found in SharePoint folder: '{day_folder_name}'"
            )

        path = f"{self._day_folder_path(day_folder_name)}/{transcript_name}"
        url  = f"{self._item_url(path)}:/content"

        resp = requests.get(url, headers=self._auth_headers())
        if resp.status_code == 404:
            raise FileNotFoundError(f"Transcript not accessible: '{transcript_name}'")
        resp.raise_for_status()

        doc       = Document(io.BytesIO(resp.content))
        lines     = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(lines)

        logger.info(
            f"Transcript read: '{transcript_name}' ({len(full_text)} chars)"
        )
        return full_text

    def read_processing_status(self, day_folder_name: str) -> Optional[dict]:
        """
        Read processing_complete.json from the output/ folder.
        Returns None if it does not exist.
        """
        path = f"{self._output_folder_path(day_folder_name)}/processing_complete.json"
        url  = f"{self._item_url(path)}:/content"
        resp = requests.get(url, headers=self._auth_headers())

        if resp.status_code == 404:
            return None
        resp.raise_for_status()
        return resp.json()

    # ------------------------------------------------------------------
    # Public: Write
    # ------------------------------------------------------------------

    def _upload_bytes(
        self,
        library_relative_path: str,
        content_bytes: bytes,
        content_type: str,
    ):
        """
        Upload raw bytes to a library-relative path.
        Graph API auto-creates the output/ folder on first upload.
        Overwrites silently if the file already exists.
        """
        drive_id = self._get_drive_id()
        url = (
            f"{GRAPH_BASE}/drives/{drive_id}"
            f"/root:/{library_relative_path}:/content"
        )
        resp = requests.put(
            url,
            headers={
                "Authorization": f"Bearer {self._get_token()}",
                "Content-Type": content_type,
            },
            data=content_bytes,
        )
        resp.raise_for_status()
        logger.info(f"Uploaded: {library_relative_path}")

    def upload_docx(self, day_folder_name: str, file_name: str, paragraphs: list):
        """
        Build a .docx in memory from a list of paragraph strings and upload
        it to the Day folder's output/ directory.

        Args:
            day_folder_name : e.g. "Day 2 - AI training - Pavan"
            file_name       : e.g. "summary.docx" or "qa.docx"
            paragraphs      : list of strings, each becomes one paragraph
        """
        doc = Document()
        for para in paragraphs:
            doc.add_paragraph(para)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        path = f"{self._output_folder_path(day_folder_name)}/{file_name}"
        self._upload_bytes(
            path,
            buffer.read(),
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document",
        )

    def upload_processing_status(self, day_folder_name: str, phase1: bool, phase2: bool):
        """Write / overwrite processing_complete.json in output/."""
        payload = {
            "phase1Completed": phase1,
            "phase2Completed": phase2,
            "processedAt": datetime.utcnow().isoformat(),
        }
        path = f"{self._output_folder_path(day_folder_name)}/processing_complete.json"
        self._upload_bytes(
            path,
            json.dumps(payload, indent=2).encode("utf-8"),
            "application/json",
        )
    def delete_output_files(self, day_folder_name: str):
        """
        Delete all files in the output folder to force reprocessing.
        Deletes: summary.docx, qa.docx, processing_complete.json
        """
        output_folder_path = self._output_folder_path(day_folder_name)
        drive_id = self._get_drive_id()
        
        # List all files in output folder
        url = f"{GRAPH_BASE}/drives/{drive_id}/root:/{output_folder_path}:/children"
        resp = requests.get(url, headers=self._auth_headers())
        
        if resp.status_code == 404:
            # Output folder doesn't exist, nothing to delete
            logger.info(f"No output folder found for '{day_folder_name}'")
            return
        
        resp.raise_for_status()
        items = resp.json().get("value", [])
        
        # Delete each file
        for item in items:
            if "file" in item:  # Only delete files, not folders
                item_id = item["id"]
                delete_url = f"{GRAPH_BASE}/drives/{drive_id}/items/{item_id}"
                del_resp = requests.delete(delete_url, headers=self._auth_headers())
                if del_resp.status_code in (200, 204):
                    logger.info(f"Deleted {item['name']} from {day_folder_name}/output/")
                else:
                    logger.warning(f"Failed to delete {item['name']}: {del_resp.status_code}")

    # ------------------------------------------------------------------
    # Public: Read output docs (for chat queries)
    # ------------------------------------------------------------------

    def list_day_folders(self) -> list:
        """
        Return a list of all Day folder names under TRAINING_ROOT.
        Only returns folders (items without a 'file' key).
        """
        drive_id = self._get_drive_id()
        url = f"{GRAPH_BASE}/drives/{drive_id}/root:/{TRAINING_ROOT}:/children"
        resp = requests.get(url, headers=self._auth_headers())
        if resp.status_code == 404:
            return []
        resp.raise_for_status()
        items = resp.json().get("value", [])
        # Only folders, sorted by name
        folders = [i["name"] for i in items if "folder" in i]
        folders.sort()
        logger.info(f"Found {len(folders)} day folders in SharePoint")
        return folders

    def read_output_docx(self, day_folder_name: str, file_name: str) -> Optional[str]:
        """
        Read a docx from the output/ folder and return plain text.
        Returns None if the file does not exist.
        Used for reading summary.docx and qa.docx for chat queries.
        """
        path = f"{self._output_folder_path(day_folder_name)}/{file_name}"
        url = f"{self._item_url(path)}:/content"
        resp = requests.get(url, headers=self._auth_headers())
        if resp.status_code == 404:
            return None
        resp.raise_for_status()
        doc = Document(io.BytesIO(resp.content))
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return "\n".join(lines)

    def read_all_summaries(self) -> dict:
        """
        Read summary.docx from every processed Day folder.
        Returns {folder_name: summary_text} for folders that have a summary.
        Silently skips unprocessed folders.
        """
        folders = self.list_day_folders()
        summaries = {}
        for folder in folders:
            text = self.read_output_docx(folder, "summary.docx")
            if text:
                summaries[folder] = text
                logger.debug(f"Loaded summary for '{folder}'")
            else:
                logger.debug(f"No summary found for '{folder}', skipping")
        logger.info(f"Loaded summaries from {len(summaries)}/{len(folders)} folders")
        return summaries

    def read_all_qa(self) -> dict:
        """
        Read qa.docx from every processed Day folder.
        Returns {folder_name: qa_text} for folders that have a qa file.
        Silently skips unprocessed folders.
        """
        folders = self.list_day_folders()
        qa_data = {}
        for folder in folders:
            text = self.read_output_docx(folder, "qa.docx")
            if text:
                qa_data[folder] = text
                logger.debug(f"Loaded Q&A for '{folder}'")
        logger.info(f"Loaded Q&A from {len(qa_data)}/{len(folders)} folders")
        return qa_data

    def upload_audio(self, day_folder_name: str, filename: str, content_bytes: bytes) -> str:
        """
        Upload an audio .mp3 file to the output/ folder on SharePoint.
        Returns the SharePoint web URL of the uploaded file.
        """
        path = f"{self._output_folder_path(day_folder_name)}/{filename}"
        drive_id = self._get_drive_id()
        url = f"{GRAPH_BASE}/drives/{drive_id}/root:/{path}:/content"

        resp = requests.put(
            url,
            headers={
                "Authorization": f"Bearer {self._get_token()}",
                "Content-Type": "audio/mpeg",
            },
            data=content_bytes,
        )
        resp.raise_for_status()
        web_url = resp.json().get("webUrl", "")
        logger.info(f"Uploaded audio: {filename} → {web_url}")
        return web_url
    
    def download_audio(self, day_folder_name: str, filename: str) -> Optional[bytes]:
        """Download an audio file from output/ folder. Returns None if not found."""
        path = f"{self._output_folder_path(day_folder_name)}/{filename}"
        url = f"{self._item_url(path)}:/content"
        resp = requests.get(url, headers=self._auth_headers())
        if resp.status_code == 404:
            return None
        resp.raise_for_status()
        return resp.content


    def upload_audio_metadata(self, day_folder_name: str, metadata: list):
        """Upload audio_metadata.json to the output/ folder on SharePoint."""
        path = f"{self._output_folder_path(day_folder_name)}/audio_metadata.json"
        self._upload_bytes(
            path,
            json.dumps(metadata, indent=2).encode("utf-8"),
            "application/json",
        )
        logger.info(f"audio_metadata.json uploaded for '{day_folder_name}' with {len(metadata)} files")

    def upload_processing_status_with_audio(self, day_folder_name: str, audio_completed: bool):
        """
        Update processing_complete.json to include audioCompleted flag.
        Preserves existing phase1/phase2 status.
        """
        existing = self.read_processing_status(day_folder_name) or {}
        payload = {
            "phase1Completed": existing.get("phase1Completed", False),
            "phase2Completed": existing.get("phase2Completed", False),
            "audioCompleted": audio_completed,
            "processedAt": existing.get("processedAt", datetime.utcnow().isoformat()),
            "audioProcessedAt": datetime.utcnow().isoformat(),
        }
        path = f"{self._output_folder_path(day_folder_name)}/processing_complete.json"
        self._upload_bytes(
            path,
            json.dumps(payload, indent=2).encode("utf-8"),
            "application/json",
        )

    def read_audio_metadata(self, day_folder_name: str) -> list:
        """
        Read audio_metadata.json from the output/ folder.
        Returns [] if not found.
        """
        path = f"{self._output_folder_path(day_folder_name)}/audio_metadata.json"
        url = f"{self._item_url(path)}:/content"
        resp = requests.get(url, headers=self._auth_headers())
        if resp.status_code == 404:
            return []
        resp.raise_for_status()
        try:
            data = resp.json()
            return data if isinstance(data, list) else []
        except Exception as e:
            logger.warning(f"Failed to parse audio_metadata.json: {e}")
            return []


























